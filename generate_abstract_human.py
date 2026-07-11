from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_paragraph('ABSTRACT')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(24)

abstract_text = (
    "Mobile communications have exploded over the past decade, bringing with them a massive wave of unsolicited advertisements, "
    "phishing attempts, and malicious links sent directly to our phones. Older rule-based filtering systems simply cannot keep "
    "up with the fast-paced and evasive tactics modern scammers use. To solve this growing security issue, we designed and "
    "implemented a completely intelligent spam detection system powered by advanced machine learning and natural language processing. "
    "We chose a Multinomial Naive Bayes classifier as our core engine and paired it with a sophisticated text feature extractor "
    "known as TF-IDF. Cybersecurity datasets almost always suffer from a severe imbalance because legitimate messages heavily outnumber "
    "spam threats. We successfully countered this problem by applying the Synthetic Minority Over-sampling Technique to synthetically "
    "balance our training data. We then trained our model on the UCI SMS Spam Collection dataset after running the text through a "
    "strict five-stage cleaning pipeline that handled noise removal, tokenization, and stopword elimination. We subsequently "
    "deployed the final model as a modern web application featuring a fast Python backend and an interactive React dashboard. "
    "Our evaluation revealed that the enhanced model achieved an impressive 96.32% overall accuracy alongside a 95.30% recall rate. "
    "These results prove that the system is highly effective at identifying disguised threats while keeping false alarms to a bare "
    "minimum. Ultimately, we created a real-time and highly scalable tool that helps network operators and everyday users intercept "
    "malicious text messages before they can do any harm."
)

body = doc.add_paragraph(abstract_text)
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
for run in body.runs:
    run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\5_Abstract_Humanized.docx')
print('Abstract page created!')
