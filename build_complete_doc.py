import re, os, sys, copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DOCX_DIR = r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx'

chapters_paths = [
    r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\INTELLIGENT SPAM DETECTION SYSTEM FOR MESSAGING PLATFORMS 1-5.docx',
    os.path.join(DOCX_DIR, 'INTELLIGENT SPAM DETECTION SYSTEM FOR MESSAGING PLATFORMS 1-5.docx')
]
CHAPTERS_FILE = next((p for p in chapters_paths if os.path.exists(p)), None)
if not CHAPTERS_FILE:
    raise FileNotFoundError("Cannot find chapters document")

OUTPUT = os.path.join(DOCX_DIR, 'Complete_Project_Document.docx')

# =====================================================================
# HELPER FUNCTIONS
# =====================================================================
def mk(tag, attribs=None, text=None):
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    if text is not None:
        el.text = text
    return el

def make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    r.append(mk('w:br', {'w:type': 'page'}))
    p.append(r)
    return p

def make_section_break_para(fmt, start, suppress_first_page=False):
    """Creates a paragraph that ends a section with specific page numbering."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    sectPr.append(mk('w:type', {'w:val': 'nextPage'}))
    sectPr.append(mk('w:pgNumType', {'w:fmt': fmt, 'w:start': str(start)}))
    if suppress_first_page:
        sectPr.append(OxmlElement('w:titlePg'))
    pPr.append(sectPr)
    p.append(pPr)
    return p

def make_centered_bold(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pPr.append(mk('w:jc', {'w:val': 'center'}))
    pPr.append(mk('w:spacing', {'w:before': '240', 'w:after': '240'}))
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_text_para(text, double_space=False, justify=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if double_space:
        pPr.append(mk('w:spacing', {'w:line': '480', 'w:lineRule': 'auto'}))
    if justify:
        pPr.append(mk('w:jc', {'w:val': 'both'}))
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_toc_field():
    toc_p = OxmlElement('w:p')
    r1 = OxmlElement('w:r')
    r1.append(mk('w:fldChar', {'w:fldCharType': 'begin', 'w:dirty': 'true'}))
    toc_p.append(r1)
    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    toc_p.append(r2)
    r3 = OxmlElement('w:r')
    r3.append(mk('w:fldChar', {'w:fldCharType': 'separate'}))
    toc_p.append(r3)
    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.text = '[Open in Word > Select All (Ctrl+A) > Press F9 to generate TOC automatically]'
    r4.append(t4)
    toc_p.append(r4)
    r5 = OxmlElement('w:r')
    r5.append(mk('w:fldChar', {'w:fldCharType': 'end'}))
    toc_p.append(r5)
    return toc_p

def get_body_elements(filepath):
    """Get all body elements from a docx file (excluding final sectPr)."""
    if not os.path.exists(filepath):
        print(f"  WARNING: File not found: {os.path.basename(filepath)}")
        return []
    doc = Document(filepath)
    elements = []
    for el in doc.element.body:
        if el.tag == qn('w:sectPr'):
            continue
        elements.append(copy.deepcopy(el))
    return elements

def set_page_number_footer(section, linked=False):
    """Set a centered page number in the footer of a section."""
    footer = section.footer
    footer.is_linked_to_previous = linked
    if linked:
        return
    # Remove existing paragraphs
    for p in list(footer.paragraphs):
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el)
    # Add new centered page number paragraph
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fc1)
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)

# =====================================================================
# BUILD THE DOCUMENT
# =====================================================================
print("Loading chapters document...")
main_doc = Document(CHAPTERS_FILE)

# Apply heading styles to chapter content
chapter_re = re.compile(r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE)$', re.IGNORECASE)
h3_re = re.compile(r'^\d+\.\d+\.\d+\b')
h2_re = re.compile(r'^\d+\.\d+\b')

for para in main_doc.paragraphs:
    t = para.text.strip()
    if not t:
        continue
    if chapter_re.match(t):
        para.style = 'Heading 1'
    elif h3_re.match(t):
        para.style = 'Heading 3'
    elif h2_re.match(t):
        para.style = 'Heading 2'
print("  Heading styles applied.")

body = main_doc.element.body
first_el = body[0]  # First element of Chapter 1 content

# =====================================================================
# Build insert list (all elements to prepend, in correct order)
# =====================================================================
insert_list = []

# ------ 1. COVER PAGE (no page number) ------
print("Adding cover page...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '1_Title_Page.docx')))
# Section break after cover page: this ends section 1 (no page numbers)
# We use titlePg to suppress the FIRST page footer of section 2
# Actually, the cover page section itself should have no footer.
# We end section 1 here with no page numbering. Section 2 starts with roman ii.
insert_list.append(make_section_break_para('lowerRoman', 2, suppress_first_page=False))

# ------ 2. DECLARATION (page ii) ------
print("Adding declaration...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '2b_Declaration.docx')))
insert_list.append(make_page_break_para())

# ------ 3. CERTIFICATION (page iii) ------
print("Adding certification...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '2_Certification_Final.docx')))
insert_list.append(make_page_break_para())

# ------ 4. DEDICATION (page iv) ------
print("Adding dedication...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '3_Dedication.docx')))
insert_list.append(make_page_break_para())

# ------ 5. ACKNOWLEDGEMENTS (page v) ------
print("Adding acknowledgements...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '4_Acknowledgement.docx')))
insert_list.append(make_page_break_para())

# ------ 6. ABSTRACT (page vi) ------
print("Adding abstract...")
insert_list.extend(get_body_elements(os.path.join(DOCX_DIR, '5_Abstract_Humanized.docx')))
insert_list.append(make_page_break_para())

# ------ 7. TABLE OF CONTENTS ------
print("Adding Table of Contents...")
insert_list.append(make_centered_bold('TABLE OF CONTENTS'))
insert_list.append(make_toc_field())
insert_list.append(make_page_break_para())

# ------ 8. LIST OF TABLES ------
print("Adding List of Tables...")
insert_list.append(make_centered_bold('LIST OF TABLES'))
insert_list.append(OxmlElement('w:p'))
for item in [
    'Table 4.1: Dataset Distribution',
    'Table 4.2: Confusion Matrix (Baseline Model vs. SMOTE Model)',
    'Table 4.3: Performance Metric Comparison',
]:
    insert_list.append(make_text_para(item, double_space=True))
insert_list.append(make_page_break_para())

# ------ 9. LIST OF FIGURES ------
print("Adding List of Figures...")
insert_list.append(make_centered_bold('LIST OF FIGURES'))
insert_list.append(OxmlElement('w:p'))
for item in [
    'Figure 3.1: High-Level System Architecture Diagram',
    'Figure 3.2: Machine Learning Pipeline Flowchart',
    'Figure 4.1: Login Page',
    'Figure 4.2: Dashboard Page (Scanner Interface)',
    'Figure 4.3: Dashboard - Spam Detection Result',
    'Figure 4.4: Dashboard - Safe Message Result',
    'Figure 4.5: Analytics Page',
    'Figure 4.6: Threat Feeds Page',
    'Figure 4.7: Settings Page',
]:
    insert_list.append(make_text_para(item, double_space=True))

# Section break before Chapter 1: ends roman numeral section, starts Arabic from 1
insert_list.append(make_section_break_para('decimal', 1))

# Insert all elements before Chapter 1 (reversed to maintain order)
for el in reversed(insert_list):
    first_el.addprevious(el)

print("All pages inserted in correct order.")

# =====================================================================
# FINAL SECTION: Arabic numbering from 1 (for chapters)
# =====================================================================
body_sectPr = main_doc.element.body.find(qn('w:sectPr'))
if body_sectPr is None:
    body_sectPr = OxmlElement('w:sectPr')
    main_doc.element.body.append(body_sectPr)

existing_pg = body_sectPr.find(qn('w:pgNumType'))
if existing_pg is not None:
    body_sectPr.remove(existing_pg)
body_sectPr.append(mk('w:pgNumType', {'w:fmt': 'decimal', 'w:start': '1'}))

# Add page number footers to ALL sections
for i, section in enumerate(main_doc.sections):
    # Section 0 = cover page: no footer (linked but empty)
    if i == 0:
        section.footer.is_linked_to_previous = False
        for p in list(section.footer.paragraphs):
            el = p._element
            if el.getparent() is not None:
                el.getparent().remove(el)
        section.footer.add_paragraph()  # empty footer = no page number
    else:
        set_page_number_footer(section)

print("Page numbering configured.")

# =====================================================================
# SAVE
# =====================================================================
main_doc.save(OUTPUT)

print("\n" + "="*60)
print("SUCCESS! File saved:")
print("  Complete_Project_Document.docx")
print("="*60)
print("\nFINAL STEP (takes 5 seconds in Word):")
print("1. Open Complete_Project_Document.docx")
print("2. Press Ctrl+A  (select all)")
print("3. Press F9      (update all fields)")
print("4. Choose: 'Update entire table'")
print("Your Table of Contents will fill in automatically!")
print("="*60)
