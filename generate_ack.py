from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph('ACKNOWLEDGEMENTS')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(24)

paragraphs = [
    "I give all glory, honor, and gratitude to Almighty God for His unfailing love, mercy, protection, and guidance throughout my academic journey and during the completion of this research work.",
    
    "My sincere appreciation goes to my project supervisor, Prof. Moses Aregbesola, for the valuable guidance, constructive criticisms, encouragement, and professional support provided throughout the course of this study. Your patience and commitment greatly contributed to the successful completion of this work.",
    
    "I also express my profound gratitude to all the lecturers in the Department of Cyber Security for their dedication to teaching and for imparting the knowledge and skills that have shaped my academic development.",
    
    "Special thanks go to my parents, Mr. and Mrs. Ogundiran, my sister, OluwaDamilola Ojo, my brother, Ogundiran Ayomide, and my cousin, Oludiran Oluwatimileyin, for their unwavering support, prayers, encouragement, and sacrifices throughout my education. Their love and belief in me have been a constant source of motivation.",
    
    "I want to also appreciate my bosses, Mr. Saheed Afolabi, Oluwatooni Oloye, and Omowaye Sylvester for their strong support and assistance throughout this journey.",
    
    "I want to appreciate my friends, Clement, Damilare, Dave, Jamal-Deen, David, Divine, Olubo, Jesse, Ashiru, and Jerry, for their assistance in all my academic endeavors.",
    
    "To my brothers, Kolade Mubeen, Adewale Ayomide, Durojaiye Daniel, Olorunfemi Emmanuel, Olukoya Olakunle, Yusuf Abdulmalik, and Steven, for always being there despite the bad times.",
    
    "To my sisters, Ayoade Precious, Adeyemo Ibukunmi, and Alotunsin Dorcas, and my love, Lawal Zainab Temiloluwa, for the continuous love and support all throughout this journey.",
    
    "I acknowledge all authors, researchers, and institutions whose scholarly works were consulted and referenced in this study. Their contributions provided valuable insights that enriched this research.",
    
    "Last but not the least, I want to thank me. I want to thank me for believing in me. I want to thank me for doing all this hardwork. I want to thank me for having no days off. I want to thank me for never quitting. I want to thank me for always being a giver, trying to give more than I receive. I want to thank me for trying to do more right than wrong. I want to thank me for just being me at all times. Ogundiran Caleb, you are special.",
    
    "May God bless everyone who contributed to the success of this work."
]

for text in paragraphs:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\4_Acknowledgement.docx')
print('Acknowledgement page created!')
