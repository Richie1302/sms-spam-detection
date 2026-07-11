import re, os, sys
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Paths ---
input_paths = [
    r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\INTELLIGENT SPAM DETECTION SYSTEM FOR MESSAGING PLATFORMS 1-5.docx',
    r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\INTELLIGENT SPAM DETECTION SYSTEM FOR MESSAGING PLATFORMS 1-5.docx'
]
INPUT = None
for p in input_paths:
    if os.path.exists(p):
        INPUT = p
        break
if not INPUT:
    raise FileNotFoundError("Could not find the source document. Check the file path.")

OUTPUT = r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\Full_Document_With_TOC.docx'

print("Source document found and loaded.")
doc = Document(INPUT)

# --- Step 1: Apply Heading Styles ---
chapter_re = re.compile(r'^CHAPTER\s+(ONE|TWO|THREE|FOUR|FIVE)$', re.IGNORECASE)
h3_re = re.compile(r'^\d+\.\d+\.\d+\b')
h2_re = re.compile(r'^\d+\.\d+\b')

for para in doc.paragraphs:
    t = para.text.strip()
    if not t:
        continue
    if chapter_re.match(t):
        para.style = 'Heading 1'
    elif h3_re.match(t):
        para.style = 'Heading 3'
    elif h2_re.match(t):
        para.style = 'Heading 2'

print("Heading styles applied.")

# --- Helper: make XML element ---
def mk(tag, attribs=None, text=None):
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    if text is not None:
        el.text = text
    return el

# --- Step 2: Insert TOC field + LOT + LOF at start ---
body = doc.element.body
first_elem = body[0]

# List of Figures content
lof_items = [
    'Figure 3.1: High-Level System Architecture Diagram',
    'Figure 3.2: Machine Learning Pipeline Flowchart',
    'Figure 4.1: Login Interface',
    'Figure 4.2: Dashboard Page (Scanner Interface)',
    'Figure 4.3: Dashboard - Spam Detection Result',
    'Figure 4.4: Dashboard - Safe Message Result',
    'Figure 4.5: Analytics Page',
    'Figure 4.6: Threat Feeds Page',
    'Figure 4.7: Settings Page',
]

# List of Tables content
lot_items = [
    'Table 4.1: Dataset Distribution',
    'Table 4.2: Confusion Matrix (Baseline Model vs. SMOTE Model)',
    'Table 4.3: Performance Metric Comparison',
]

def make_centered_bold_para(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = mk('w:jc', {'w:val': 'center'})
    pPr.append(jc)
    spacing = mk('w:spacing', {'w:before': '240', 'w:after': '240'})
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_list_item(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    spacing = mk('w:spacing', {'w:line': '480', 'w:lineRule': 'auto'})
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_page_break():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = mk('w:br', {'w:type': 'page'})
    r.append(br)
    p.append(r)
    return p

def make_blank():
    return OxmlElement('w:p')

# Build elements to prepend (in reverse order since we use addprevious)
# Order we want: TOC title, TOC field, page break, LOT title, LOT items, page break, LOF title, LOF items, page break

elements_to_insert = []

# TOC title
elements_to_insert.append(make_centered_bold_para('TABLE OF CONTENTS'))

# TOC field
toc_p = OxmlElement('w:p')
r1 = OxmlElement('w:r')
r1.append(mk('w:fldChar', {'w:fldCharType': 'begin', 'w:dirty': 'true'}))
toc_p.append(r1)
r2 = OxmlElement('w:r')
instr = OxmlElement('w:instrText')
instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
instr.text = ' TOC \\o "1-3" \\h \\z \\u '
r2.append(instr)
toc_p.append(r2)
r3 = OxmlElement('w:r')
r3.append(mk('w:fldChar', {'w:fldCharType': 'separate'}))
toc_p.append(r3)
r4 = OxmlElement('w:r')
t4 = OxmlElement('w:t')
t4.text = '[Right-click this line in Word and select "Update Field" to generate the TOC]'
r4.append(t4)
toc_p.append(r4)
r5 = OxmlElement('w:r')
r5.append(mk('w:fldChar', {'w:fldCharType': 'end'}))
toc_p.append(r5)
elements_to_insert.append(toc_p)

# Page break after TOC
elements_to_insert.append(make_page_break())

# LIST OF TABLES
elements_to_insert.append(make_centered_bold_para('LIST OF TABLES'))
elements_to_insert.append(make_blank())
for item in lot_items:
    elements_to_insert.append(make_list_item(item))
elements_to_insert.append(make_page_break())

# LIST OF FIGURES
elements_to_insert.append(make_centered_bold_para('LIST OF FIGURES'))
elements_to_insert.append(make_blank())
for item in lof_items:
    elements_to_insert.append(make_list_item(item))
elements_to_insert.append(make_page_break())

# Insert all before the first element (in reverse so order is preserved)
for el in reversed(elements_to_insert):
    first_elem.addprevious(el)

print("TOC, LOT and LOF inserted at beginning.")

# --- Step 3: Page Numbering ---
# Find Chapter 1 paragraph
chapter1_el = None
for para in doc.paragraphs:
    if chapter_re.match(para.text.strip()):
        chapter1_el = para._element
        break

if chapter1_el:
    # Get the paragraph just before Chapter 1
    prev = chapter1_el.getprevious()
    while prev is not None and prev.tag != qn('w:p'):
        prev = prev.getprevious()

    if prev is not None:
        pPr = prev.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            prev.insert(0, pPr)
        # Remove existing sectPr if any
        existing = pPr.find(qn('w:sectPr'))
        if existing is not None:
            pPr.remove(existing)

        # Create section properties for preliminary pages (roman numerals from ii)
        sectPr = OxmlElement('w:sectPr')
        sectPr.append(mk('w:pgNumType', {'w:fmt': 'lowerRoman', 'w:start': '2'}))
        # Suppress first page number (cover page)
        sectPr.append(OxmlElement('w:titlePg'))
        # Page size (Letter)
        sectPr.append(mk('w:pgSz', {'w:w': '12240', 'w:h': '15840'}))
        sectPr.append(mk('w:pgMar', {
            'w:top': '1440', 'w:right': '1440',
            'w:bottom': '1440', 'w:left': '1800',
            'w:header': '720', 'w:footer': '720'
        }))

        # Footer with roman page number for preliminary section
        ftr_ref = mk('w:footerReference', {'w:type': 'default'})
        # We'll add footer via section footer
        pPr.append(sectPr)
        print("Preliminary section break (roman numerals) added before Chapter 1.")

# Final section (chapters) - Arabic from 1
body_sectPr = body.find(qn('w:sectPr'))
if body_sectPr is None:
    body_sectPr = OxmlElement('w:sectPr')
    body.append(body_sectPr)

existing_pg = body_sectPr.find(qn('w:pgNumType'))
if existing_pg is not None:
    body_sectPr.remove(existing_pg)

body_sectPr.append(mk('w:pgNumType', {'w:fmt': 'decimal', 'w:start': '1'}))
print("Chapter section (Arabic from 1) configured.")

# --- Step 4: Add page number to footers using section footer API ---
from docx.oxml import OxmlElement as OE

def set_footer_page_number(section, fmt):
    """Add centered page number to footer"""
    section.footer.is_linked_to_previous = False
    footer = section.footer
    # Clear existing content
    for p in footer.paragraphs:
        el = p._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    # Add paragraph with page number
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # PAGE field
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fc1)
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)
    fc2 = OxmlElement('w:fldChar')
    fc2.set(qn('w:fldCharType'), 'end')
    run._r.append(fc2)

# Apply footers to all sections
for i, section in enumerate(doc.sections):
    set_footer_page_number(section, 'page')

print("Page number footers added to all sections.")

# --- Save ---
doc.save(OUTPUT)
print(f"\nDone! Saved to:\n{OUTPUT}")
print("\n" + "="*60)
print("IMPORTANT - Final steps in Microsoft Word:")
print("1. Open Full_Document_With_TOC.docx in Word")
print("2. Press Ctrl+A to select all text")
print("3. Press F9 to update all fields (this generates the TOC)")
print("4. If prompted, choose 'Update entire table'")
print("="*60)
