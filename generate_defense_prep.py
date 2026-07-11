from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    return p

def body(text, bold_prefix=None, double=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if double:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.font.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p

def tip_box(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("💡 TIP: " + text)
    r.font.italic = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    return p

def separator():
    p = doc.add_paragraph("─" * 70)
    p.paragraph_format.space_after = Pt(4)

# ============================================================
# COVER / TITLE
# ============================================================
t = doc.add_paragraph("DEFENSE PREPARATION GUIDE")
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.runs[0]; r.font.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'
r.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

t2 = doc.add_paragraph("Intelligent SMS Spam Detection System")
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.runs[0]; r2.font.bold = True; r2.font.size = Pt(13); r2.font.name = 'Times New Roman'

doc.add_paragraph()

# ============================================================
# SECTION 0: QUICK IDENTITY (2 MINUTES INTRO)
# ============================================================
heading("SECTION 0 — YOUR 2-MINUTE OPENING SPEECH")
body('This is what you say the moment they ask you to introduce yourself and your project. Memorise this word for word.')
separator()

body(
    'Good morning/afternoon. My name is Ogundiran Caleb Ayodeji, '
    'with Matriculation Number 22/9789, from the Department of Cyber Security, '
    'Caleb University. My project is titled "Design and Implementation of an '
    'Intelligent SMS Spam Detection System." I chose this project because SMS spam '
    'and smishing — which is SMS phishing — has become a serious cybersecurity '
    'threat in Nigeria and globally. Traditional spam filters could not keep up with '
    'the evolving tactics of scammers, so I decided to build a smarter, '
    'machine-learning-powered system that can detect these threats in real time. '
    'The system is now fully deployed as a live web application.'
)
tip_box("Smile when you say this. Confidence is 50% of your grade.")

# ============================================================
# SECTION 1: NAME & BASIC INFO
# ============================================================
heading("SECTION 1 — BASIC INFORMATION")
separator()
body("Ogundiran Caleb Ayodeji", bold_prefix="Full Name:")
body("22/9789", bold_prefix="Matric Number:")
body("Department of Cyber Security", bold_prefix="Department:")
body("Caleb University, Imota, Ikorodu, Lagos", bold_prefix="University:")
body("Design and Implementation of an Intelligent SMS Spam Detection System", bold_prefix="Project Title:")
body("Prof. Moses Aregbesola", bold_prefix="Supervisor:")
body("June, 2026", bold_prefix="Date:")

# ============================================================
# SECTION 2: WHY I CHOSE THIS PROJECT
# ============================================================
heading("SECTION 2 — WHY I CHOSE THIS PROJECT")
body('If they ask: "Why did you choose this topic?" — Say this:')
separator()
body(
    'I chose this project because SMS spam is not just annoying — it is dangerous. '
    'In Nigeria, millions of people receive fraudulent text messages every day that '
    'try to steal their bank details, OTP codes, and personal information. These '
    'attacks are called "smishing" (SMS phishing). Existing filters were either too '
    'simple or kept blocking legitimate messages. I wanted to build something '
    'intelligent — a system that actually learns the difference between a real '
    'message and a threat, and gets better over time. As a Cyber Security student, '
    'this was the most practical and impactful project I could work on.'
)
tip_box("If they push: 'Why not email spam?' — Answer: SMS is more personal, more trusted, and more dangerous because people are less suspicious of text messages than emails.")

# ============================================================
# SECTION 3: GENERAL INTRODUCTION (IN SIMPLE TERMS)
# ============================================================
heading("SECTION 3 — GENERAL INTRODUCTION")
body('What is your project basically about? Explain it like you are talking to a non-technical person.')
separator()
body(
    'My project is a smart system that reads an SMS message and tells you — within '
    'seconds — whether it is a legitimate message (called "ham") or a spam/scam '
    'message (called "spam"). Think of it like a security guard at a gate. Every '
    'message that comes in goes through the guard. The guard has been trained on '
    'thousands of real and fake messages, so he knows the patterns scammers use. '
    'If the message looks like a threat, the system raises a red alert. If it is '
    'safe, it gives a green light. The system was built using Artificial Intelligence '
    '(Machine Learning), and it is accessible through a website that anyone can open '
    'on their browser and use instantly.'
)
tip_box("Remember: the live website is at https://sms-spam-detection-weld.vercel.app/ — mention this! It shows your project is fully deployed, not just a theory.")

# ============================================================
# SECTION 4: PROBLEM STATEMENT
# ============================================================
heading("SECTION 4 — PROBLEM STATEMENT")
body('If they ask: "What problem are you solving?" — Say this:')
separator()
body(
    'The problem I am solving is the inadequacy of existing SMS spam detection '
    'methods in the face of modern, evolving cyber threats. Traditional spam filters '
    'rely on simple keyword blacklists and rigid rules, which scammers easily bypass '
    'by changing a few words. More critically, these systems suffer from a class '
    'imbalance problem — since legitimate messages are far more common than spam, '
    'the filter becomes biased and either misses actual spam or wrongly flags real '
    'messages as spam. This research addresses this problem by designing a machine '
    'learning-based detection system that learns the statistical patterns of spam '
    'and uses advanced techniques to handle the imbalance in training data.'
)
tip_box("In simple terms: 'The old systems were too rigid and kept making mistakes. I built a smarter one that learns and adapts.'")

# ============================================================
# SECTION 5: AIM & OBJECTIVES
# ============================================================
heading("SECTION 5 — AIM AND OBJECTIVES")
body('The AIM is the big goal. The OBJECTIVES are the steps to achieve it.')
separator()
body(
    'To design, implement, and deploy an intelligent SMS spam detection system '
    'using machine learning and Natural Language Processing techniques.',
    bold_prefix="AIM:"
)
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("OBJECTIVES:")
r.font.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

objectives = [
    "To collect and preprocess a real-world SMS dataset for model training.",
    "To implement a TF-IDF feature extraction pipeline for converting raw text into numerical features the machine can understand.",
    "To train a Multinomial Naive Bayes classifier to distinguish spam from legitimate messages.",
    "To apply SMOTE (Synthetic Minority Over-sampling Technique) to handle the class imbalance problem in the dataset.",
    "To evaluate the model's performance using metrics such as Accuracy, Precision, Recall, and the Matthews Correlation Coefficient.",
    "To deploy the trained model as a real-time web application accessible via browser."
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(obj)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ============================================================
# SECTION 6: LITERATURE REVIEW (SIMPLE VERSION WITH REFS)
# ============================================================
heading("SECTION 6 — LITERATURE REVIEW (WITH REFERENCES)")
body('If they ask about related works or what other researchers have done, say this:')
separator()
body(
    'Several researchers have worked on SMS spam detection. Early works like Almeida '
    'et al. (2011) introduced the UCI SMS Spam Collection dataset — which is the '
    'very dataset I used — and demonstrated that machine learning models could '
    'significantly outperform manual filtering. Uysal et al. (2013) showed that '
    'preprocessing steps like removing stopwords greatly improve classifier accuracy. '
    'More recently, researchers like Saidat et al. (2024) confirmed that Multinomial '
    'Naive Bayes with TF-IDF remains one of the most effective and computationally '
    'efficient approaches for SMS spam detection. Chen et al. (2024) further '
    'demonstrated the value of combining TF-IDF with classical classifiers for '
    'real-world deployment. The gap identified in the literature — and the gap my '
    'research addresses — is the consistent neglect of the class imbalance problem. '
    'Faisal et al. (2023) showed that applying SMOTE before training dramatically '
    'improves recall for the minority class (spam), which is critical for a '
    'cybersecurity application where missing a spam message is far more dangerous '
    'than a false alarm.'
)
tip_box("Key papers to remember: Almeida et al. (2011) — the dataset. Saidat et al. (2024) — MNB + TF-IDF. Faisal et al. (2023) — SMOTE for class imbalance.")

# ============================================================
# SECTION 7: METHODOLOGY
# ============================================================
heading("SECTION 7 — METHODOLOGY & DATASET")
body('How did you build this system? Walk them through it step by step.')
separator()

body("UCI SMS Spam Collection Dataset — 5,572 real SMS messages (4,825 ham + 747 spam). This is a well-known, publicly available dataset used in cybersecurity research.", bold_prefix="DATASET:")
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("STEP-BY-STEP PROCESS:")
r.font.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

steps = [
    ("Step 1 — Data Collection:", "Downloaded the UCI SMS Spam Collection Dataset containing 5,572 text messages already labelled as spam or ham."),
    ("Step 2 — Text Preprocessing:", "Cleaned the raw text by converting to lowercase, removing punctuation, numbers, and special characters, tokenizing the sentences into individual words, and removing stopwords (common words like 'the', 'is', 'and' that carry no meaning for classification)."),
    ("Step 3 — Feature Extraction (TF-IDF):", "Converted the cleaned text into numbers using TF-IDF (Term Frequency-Inverse Document Frequency). This gives high scores to words that appear often in spam messages but rarely in normal messages — like 'FREE', 'WIN', 'CLAIM', 'PRIZE'."),
    ("Step 4 — Handling Class Imbalance (SMOTE):", "The dataset had far more ham messages than spam. Without fixing this, the model would just learn to call everything 'ham' and be technically accurate but useless. SMOTE created artificial but realistic spam samples to balance the training data."),
    ("Step 5 — Model Training (Multinomial Naive Bayes):", "Trained a Multinomial Naive Bayes (MNB) classifier. This algorithm uses probability — it calculates the likelihood that a message is spam based on the words it contains. It is fast, interpretable, and proven effective for text classification."),
    ("Step 6 — Evaluation:", "Tested the model using Accuracy, Precision, Recall, F1-Score, and Matthews Correlation Coefficient (MCC)."),
    ("Step 7 — Deployment:", "Built a FastAPI (Python) backend to serve the model and a React.js frontend for the user interface. The live website is at: https://sms-spam-detection-weld.vercel.app/"),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(title + " ")
    r1.font.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(desc)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

# ============================================================
# SECTION 8: LANGUAGES & TOOLS
# ============================================================
heading("SECTION 8 — LANGUAGES AND TOOLS USED")
separator()
tools = [
    ("Python 3.12", "Main programming language for machine learning and backend development."),
    ("Scikit-Learn", "Machine learning library used to build the Naive Bayes classifier and apply SMOTE."),
    ("NLTK (Natural Language Toolkit)", "Used for text preprocessing — tokenization, stopword removal."),
    ("TF-IDF Vectorizer", "Used for converting text into numerical features."),
    ("FastAPI", "A fast Python web framework used to build the REST API backend that serves the ML model."),
    ("React.js", "JavaScript framework used to build the frontend user interface (the website dashboard)."),
    ("Vercel", "Cloud platform used to host and deploy the frontend website."),
    ("Render", "Cloud platform used to deploy the backend API and ML model."),
    ("GitHub", "Used for version control and storing the source code (public repository)."),
    ("VS Code", "The code editor used for development."),
]
for tool, desc in tools:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"• {tool}: ")
    r1.font.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(desc)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

# ============================================================
# SECTION 9: RESULTS
# ============================================================
heading("SECTION 9 — RESULTS (WHAT DID YOU ACHIEVE?)")
body('If they ask about your results or findings, say this:')
separator()
body(
    'The SMOTE-enhanced Multinomial Naive Bayes model achieved the following results:'
)
metrics = [
    ("Accuracy:", "96.32% — The model correctly classified 96 out of every 100 messages."),
    ("Precision:", "89.15% — When the model says a message is spam, it is right almost 9 out of 10 times."),
    ("Recall:", "95.30% — The model catches 95% of all spam messages that exist in the dataset."),
    ("F1-Score:", "92.12% — A balanced measure combining both precision and recall."),
    ("MCC:", "0.917 — This is near-perfect on a scale of -1 to +1. It proves the model works well even on imbalanced data."),
    ("False Positive Rate:", "1.8% — Only 1.8% of real/legitimate messages were incorrectly flagged as spam."),
]
for label, val in metrics:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r1 = p.add_run(f"  {label} ")
    r1.font.bold = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(val)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

body(
    'Compared to the baseline model (trained WITHOUT SMOTE), the SMOTE version improved '
    'recall from 78% to 95.30% — a massive 17-percentage-point improvement in catching '
    'spam, which is the most critical metric for a cybersecurity application.'
)
tip_box("If they ask why recall matters more than accuracy: 'In cybersecurity, missing a spam message is far more dangerous than flagging a real message. Recall tells us how many threats we actually caught.'")

# ============================================================
# SECTION 10: FUTURE PLANS
# ============================================================
heading("SECTION 10 — FUTURE PLANS FOR THE PROJECT")
body('If they ask: "What would you improve if you had more time?"')
separator()
future = [
    "Extend the system to support Yoruba, Hausa, and Pidgin English SMS messages, since most Nigerian SMS scams are written in local dialects.",
    "Implement deep learning models (LSTM or BERT) which can understand context better than Naive Bayes for even higher accuracy.",
    "Integrate the system directly with mobile network operators' SMS gateway so filtering happens before the message even reaches the user's phone.",
    "Add a real-time threat intelligence feed that automatically updates the model with new scam patterns as they emerge.",
    "Build a mobile app version so users can forward suspicious messages for instant scanning.",
]
for item in future:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run(item)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ============================================================
# SECTION 11: LIKELY DEFENSE QUESTIONS & ANSWERS
# ============================================================
heading("SECTION 11 — LIKELY QUESTIONS & SMART ANSWERS")
separator()

qas = [
    (
        "Q1: What is the difference between spam and ham?",
        "Spam is any unsolicited, unwanted, or malicious text message — like scam links, fake prizes, or phishing attempts. Ham is a legitimate message — like a text from your bank, your friend, or a delivery service. These are the two classes my model classifies every message into."
    ),
    (
        "Q2: Why did you choose Naive Bayes and not a Neural Network?",
        "Naive Bayes is computationally efficient, interpretable, and proven highly effective for text classification tasks. Neural networks require much more data and computing power to train, and for SMS spam detection — which is a relatively simple binary classification problem — Naive Bayes gives excellent results with far less complexity. My results of 96.32% accuracy confirm this choice was justified."
    ),
    (
        "Q3: What is TF-IDF and why did you use it?",
        "TF-IDF stands for Term Frequency–Inverse Document Frequency. It converts raw text into numbers that the machine learning model can process. 'Term Frequency' measures how often a word appears in a message. 'Inverse Document Frequency' reduces the weight of words that appear in almost every message — like 'the' or 'is' — and increases the weight of rare but important words like 'FREE', 'WIN', or 'CLAIM'. This helps the model focus on the words that actually distinguish spam from legitimate messages."
    ),
    (
        "Q4: What is SMOTE and why was it necessary?",
        "SMOTE stands for Synthetic Minority Over-sampling Technique. In my dataset, there were 4,825 legitimate messages but only 747 spam messages. This imbalance means a model could score 87% accuracy just by calling everything 'ham' — but that model would be completely useless because it would miss all the actual spam. SMOTE solved this by generating synthetic but realistic spam samples to balance the training data, which dramatically improved my recall from 78% to 95.30%."
    ),
    (
        "Q5: What is the Matthews Correlation Coefficient?",
        "The MCC is considered the most reliable single metric for evaluating binary classifiers, especially on imbalanced datasets. It ranges from -1 (completely wrong) to +1 (perfectly correct), with 0 meaning random guessing. My model scored 0.917, which is near-perfect and proves the system is genuinely effective — not just artificially inflated by class imbalance."
    ),
    (
        "Q6: Is your system deployed? Can we see it?",
        "Yes! The system is fully deployed and live. The frontend is hosted on Vercel and the backend API is hosted on Render. The live URL is: https://sms-spam-detection-weld.vercel.app/ — You can open it right now and test it with any SMS message."
    ),
    (
        "Q7: What is the GitHub repository?",
        "The source code is publicly available at: https://github.com/Richie1302/sms-spam-detection"
    ),
    (
        "Q8: What are the limitations of your system?",
        "The main limitations are: First, the model was trained on English-language SMS messages, so its performance on Nigerian local language scams (Yoruba, Hausa, Pidgin) has not been tested. Second, very sophisticated smishing attacks that use subtle language could potentially evade the classifier. Third, the model is static — it does not automatically retrain itself when new spam patterns emerge without manual intervention."
    ),
    (
        "Q9: How does the web application work?",
        "The user visits the website and types or pastes an SMS message into the dashboard. The frontend sends the message to the FastAPI backend via an API call. The backend passes the text through the same preprocessing and TF-IDF pipeline used during training, then feeds it to the trained Multinomial Naive Bayes model. The model returns a prediction — SPAM or HAM — along with a confidence score, and the result is displayed on the screen with a visual threat indicator in under one second."
    ),
    (
        "Q10: Why Caleb University for this kind of project?",
        "Caleb University's Department of Cyber Security provided me with the theoretical foundation in network security, cryptography, and information assurance that is essential for understanding why SMS-based attacks are so dangerous. This project is a direct application of the knowledge I gained throughout my four years of study."
    ),
]

for q, a in qas:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(q)
    r.font.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p2.paragraph_format.left_indent = Inches(0.2)
    r2 = p2.add_run("Answer: " + a)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

# ============================================================
# FINAL NOTE
# ============================================================
heading("FINAL REMINDER")
separator()
body(
    'You built a REAL system that is LIVE on the internet. '
    'That alone puts you ahead of 90% of final year students. '
    'Walk into that defense room knowing that. Be calm, be clear, and be confident. '
    'You know this project better than any examiner in that room. '
    'Ogundiran Caleb Ayodeji — go and nail it!'
)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\Defense_Prep_Guide.docx')
print("Defense Prep Guide created!")
