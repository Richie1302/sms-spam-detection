from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_paragraph('TABLE OF CONTENTS')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(24)

toc_lines = [
    "Title Page",
    "Certification",
    "Declaration",
    "Dedication",
    "Acknowledgements",
    "Abstract",
    "Table of Contents",
    "List of Tables",
    "List of Figures",
    "",
    "CHAPTER ONE",
    "INTRODUCTION",
    "1.1 BACKGROUND OF THE STUDY",
    "1.2 PROBLEM STATEMENT",
    "1.3 AIM OF STUDY",
    "1.3.1 RESEARCH QUESTIONS",
    "1.4 SCOPE OF STUDY",
    "1.4.1 LIMITATIONS",
    "1.5 OBJECTIVES OF THE SYSTEM",
    "1.6 METHODOLOGY",
    "1.7 EXPECTED OUTCOME",
    "1.7.1 EXPECTED CONTRIBUTION TO CYBERSECURITY",
    "1.7.2 EXPECTED CONTRIBUTION TO MACHINE LEARNING RESEARCH METHODOLOGY",
    "1.7.3 EXPECTED CONTRIBUTION TO SOFTWARE ENGINEERING AND SYSTEM DEPLOYMENT",
    "1.7.4 ORGANIZATION OF THE PROJECT REPORT",
    "1.8 DEFINITION OF OPERATIONAL TERMS",
    "",
    "CHAPTER TWO",
    "2.1 INTRODUCTION",
    "2.2 CONCEPTUAL REVIEW",
    "2.2.1 The Short Message Service as a Communication Medium and Security Threat Vector",
    "2.2.2 Smishing: SMS Phishing as an Escalating Cybersecurity Threat",
    "2.2.3 Natural Language Processing in Automated Text Classification",
    "2.2.4 TF-IDF Vectorization and High-Dimensional Feature Spaces",
    "2.2.5 Multinomial Naive Bayes Classification and Probabilistic Modeling",
    "2.2.6 The Challenge of Class Imbalance in Machine Learning",
    "2.2.7 The Synthetic Minority Over-Sampling Technique (SMOTE)",
    "2.2.8 Comprehensive Classification Evaluation Metrics for Imbalanced Systems",
    "2.2.9 The Imperative of Web Application Deployment for Machine Learning Systems",
    "2.3 THEORETICAL FRAMEWORK",
    "2.3.1 Bayesian Probability Theory and Statistical Inference",
    "2.3.2 The Vector Space Model (VSM) in Information Retrieval",
    "2.3.3 Resampling Theory and Decision Boundary Modification",
    "2.4 EMPIRICAL REVIEW",
    "2.4.1 SMS Spam Detection: Surveys and Comprehensive Reviews",
    "2.4.2 Recent SMS Spam Detection: Methodological Approaches and Deployment Strategies",
    "2.4.3 Advanced Feature Extraction and Comparative Pipeline Studies",
    "2.4.4 Smishing and Cybersecurity-Oriented Detection Research",
    "2.4.5 Class Imbalance Strategies for Machine Learning Systems",
    "2.4.6 Evaluation Metrics: The Emergence of the MCC Standard",
    "2.5 SUMMARY OF LITERATURE AND IDENTIFIED GAPS",
    "",
    "CHAPTER THREE",
    "3.1 INTRODUCTION",
    "3.2 STUDY DESIGN",
    "3.3 STUDY POPULATION",
    "3.3.1 Population Definition",
    "3.3.2 Dataset Characteristics",
    "3.3.3 Dataset Justification",
    "3.4 DETERMINATION OF SAMPLE SIZE",
    "3.5 SAMPLING TECHNIQUE",
    "3.6 RESEARCH INSTRUMENT",
    "3.6.1 Text Preprocessing Pipeline",
    "3.6.2 Feature Extraction: TF-IDF Vectorization",
    "3.6.3 Handling Class Imbalance: SMOTE",
    "3.6.4 Classification Model: Multinomial Naive Bayes",
    "3.6.5 Software Tools and Environment",
    "3.7 METHOD AND ANALYSIS OF DATA",
    "3.7.1 Validation Strategy and Metrics",
    "3.8 SYSTEM ARCHITECTURE",
    "3.8.1 Machine Learning Pipeline Architecture",
    "3.8.2 Backend API Design (FastAPI)",
    "3.8.3 Frontend Interface Design",
    "3.8.4 System-Level Architecture Overview",
    "3.9 ETHICAL CONSIDERATION",
    "3.10 POTENTIAL VALUE OF RESULTS",
    "3.11 LIST OF INVESTIGATORS",
    "3.12 FUNDING",
    "",
    "CHAPTER FOUR",
    "4.0 Introduction",
    "4.1 System Implementation Overview",
    "4.1.1 Machine Learning Layer (Model Training and Serialization)",
    "4.1.2 Backend API Layer (FastAPI Server)",
    "4.1.3 Frontend Application Layer (React Interface)",
    "4.2 Program Flow",
    "4.3 User Interface (UI Design)",
    "4.3.1 Login Page",
    "4.3.2 Dashboard Page (Scanner Interface)",
    "4.3.3 Dashboard - Spam Detection Result",
    "4.3.4 Dashboard - Safe Message Result",
    "4.3.5 Analytics Page",
    "4.3.6 Threat Feeds Page",
    "4.3.7 Settings Page",
    "4.4 System Testing",
    "4.4.1 Spam Message Classification Test",
    "4.4.2 Legitimate Message Classification Test",
    "4.4.3 Session Logging Verification",
    "4.5 User Documentation",
    "4.5.1 System Access",
    "4.5.2 Scanning an SMS Message (Dashboard)",
    "4.5.3 Viewing Performance Analytics",
    "4.5.4 Monitoring Threat Feeds",
    "4.5.5 Configuring System Settings",
    "4.6 Result and Analysis",
    "4.6.1 Confusion Matrix Analysis",
    "4.6.2 Performance Metrics Summary",
    "4.6.3 Matthews Correlation Coefficient Analysis",
    "4.6.4 False Positive Rate Analysis and Operational Significance",
    "4.6.5 Comparison with Baseline Model",
    "",
    "CHAPTER FIVE",
    "5.1 Summary of Findings",
    "5.2 Conclusion",
    "5.3 Recommendations",
    "5.4 Contribution to Knowledge",
    "",
    "REFERENCES",
    "APPENDIX A: CODES",
    "APPENDIX B: USER'S MANUAL"
]

for line in toc_lines:
    p = doc.add_paragraph(line)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\6_Table_of_Contents.docx')
print('TOC Generated!')
