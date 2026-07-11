from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph('DECLARATION')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(12)

# Body paragraph - justified
body = doc.add_paragraph()
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
body.paragraph_format.space_after = Pt(48)

r1 = body.add_run('I hereby declare that this project titled "')
r1.font.size = Pt(12)

r2 = body.add_run('DESIGN AND IMPLEMENTATION OF AN INTELLIGENT SMS SPAM DETECTION SYSTEM')
r2.font.bold = True
r2.font.size = Pt(12)

r3 = body.add_run('" was carried out by me independently under the supervision of Prof. Aregbesola, Department of Cyber Security, Caleb University, Imota, Lagos. The work has not been previously submitted for the award of any degree and does not incorporate, without proper acknowledgement, any material previously published or written by another person. All sources of information consulted in the course of this work have been cited and referenced in accordance with the APA 7th Edition referencing standard.')
r3.font.size = Pt(12)

# Signature table
table = doc.add_table(rows=2, cols=2)

# Remove borders
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
tbl = table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement('w:tblBorders')
for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
    border = OxmlElement(f'w:{border_name}')
    border.set(qn('w:val'), 'none')
    tblBorders.append(border)
tblPr.append(tblBorders)

# Row 0: dotted lines
sig_cell = table.cell(0, 0)
sig_cell.text = '......................................'
sig_cell.paragraphs[0].runs[0].font.size = Pt(12)

date_dotted = table.cell(0, 1)
date_dotted.text = '......................................'
date_dotted.paragraphs[0].runs[0].font.size = Pt(12)

# Row 1: name bold | "Date"
name_cell = table.cell(1, 0)
name_para = name_cell.paragraphs[0]
name_run = name_para.add_run('Ogundiran Caleb Ayodeji 22/9789')
name_run.font.bold = True
name_run.font.size = Pt(12)

date_cell = table.cell(1, 1)
date_para = date_cell.paragraphs[0]
date_run = date_para.add_run('Date')
date_run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\2b_Declaration.docx')
print('Declaration page created!')
