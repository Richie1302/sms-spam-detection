from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def setup_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    return doc

def add_title(doc, text):
    title = doc.add_paragraph(text)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.bold = True
        run.font.size = Pt(12)
    title.paragraph_format.space_after = Pt(24)

# 1. Table of Contents
doc_toc = setup_doc()
add_title(doc_toc, 'TABLE OF CONTENTS')
toc_content = [
    "Title Page ...................................................................................................... i",
    "Certification ................................................................................................. ii",
    "Declaration .................................................................................................. iii",
    "Dedication ................................................................................................... iv",
    "Acknowledgements .................................................................................... v",
    "Abstract ...................................................................................................... vi",
    "Table of Contents ....................................................................................... vii",
    "List of Tables ............................................................................................. viii",
    "List of Figures ............................................................................................ ix",
    "",
    "CHAPTER ONE: INTRODUCTION",
    "1.1 Background of Study ............................................................................ 1",
    "1.2 Problem Statement ................................................................................ 3",
    "1.3 Aim of Study ........................................................................................ 4",
    "1.4 Scope of Study ..................................................................................... 5",
    "1.5 Objectives of the System ..................................................................... 5",
    "1.6 Methodology ........................................................................................ 6",
    "1.7 Expected Outcome ............................................................................... 7",
    "1.8 Definition of Operational Terms ......................................................... 8",
    "",
    "CHAPTER TWO: LITERATURE REVIEW",
    "2.1 Introduction ......................................................................................... 9",
    "2.2 Understanding Electronic Spam and Smishing .................................. 10",
    "2.3 Review of Related Works ................................................................... 14",
    "2.4 Theoretical Framework ....................................................................... 19",
    "",
    "CHAPTER THREE: RESEARCH METHODOLOGY",
    "3.1 Introduction ......................................................................................... 22",
    "3.2 Requirement Specification ................................................................... 23",
    "3.3 Method and Models Used ................................................................... 25",
    "3.4 System Development Tools ................................................................. 28",
    "3.5 Database / System Architecture Design .............................................. 30",
    "",
    "CHAPTER FOUR: IMPLEMENTATION, ANALYSIS AND TESTING",
    "4.0 Introduction ......................................................................................... 33",
    "4.1 System Overview ................................................................................ 34",
    "4.2 Program Flow ...................................................................................... 36",
    "4.3 System Testing .................................................................................... 39",
    "4.4 User Documentation ............................................................................ 42",
    "4.5 Result and Analysis ............................................................................. 44",
    "",
    "CHAPTER FIVE: SUMMARY, CONCLUSION AND RECOMMENDATION",
    "5.1 Summary ............................................................................................. 48",
    "5.2 Conclusion and Recommendation ....................................................... 50",
    "5.3 Further Works ..................................................................................... 52",
    "",
    "REFERENCES ........................................................................................ 54",
    "APPENDIX A (Codes) ............................................................................ 58",
    "APPENDIX B (User's Manual) .............................................................. 65"
]
for line in toc_content:
    p = doc_toc.add_paragraph(line)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(0)
doc_toc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\6_Table_of_Contents.docx')

# 2. List of Tables
doc_lot = setup_doc()
add_title(doc_lot, 'LIST OF TABLES')
lot_content = [
    "Table 4.1: Dataset Distribution ................................................................ 45",
    "Table 4.2: Confusion Matrix (Baseline Model vs. SMOTE Model) ....... 46",
    "Table 4.3: Performance Metric Comparison ............................................ 47"
]
for line in lot_content:
    p = doc_lot.add_paragraph(line)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
doc_lot.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\7_List_of_Tables.docx')

# 3. List of Figures
doc_lof = setup_doc()
add_title(doc_lof, 'LIST OF FIGURES')
lof_content = [
    "Figure 3.1: High-Level System Architecture Diagram ........................... 31",
    "Figure 3.2: Machine Learning Pipeline Flowchart .................................. 32",
    "Figure B.1: Login Interface ....................................................................... 66",
    "Figure B.2: Primary Dashboard Interface ................................................ 67",
    "Figure B.3: Threat Detected Result .......................................................... 68",
    "Figure B.4: Safe Message Result ............................................................. 69",
    "Figure B.5: Analytics Interface ................................................................ 70"
]
for line in lof_content:
    p = doc_lof.add_paragraph(line)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
doc_lof.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\8_List_of_Figures.docx')

print("TOC, LOT, LOF generated!")
