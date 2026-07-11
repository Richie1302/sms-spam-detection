from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

title = doc.add_paragraph('ABSTRACT')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(24)

abstract_text = (
    "The proliferation of mobile communications has led to a significant increase in Short Message Service (SMS) spam, "
    "which includes unsolicited advertisements, phishing attempts (smishing), and malicious links. Traditional rule-based "
    "filtering systems struggle to adapt to the dynamic and evasive nature of modern spam patterns. This research focuses "
    "on the design and implementation of an Intelligent SMS Spam Detection System using advanced machine learning and Natural "
    "Language Processing (NLP) techniques. The system utilizes a Multinomial Naive Bayes (MNB) classifier enhanced with "
    "Term Frequency-Inverse Document Frequency (TF-IDF) for feature extraction. To address the inherent class imbalance present "
    "in cybersecurity datasets—where legitimate messages (ham) significantly outnumber spam—the Synthetic Minority Over-sampling "
    "Technique (SMOTE) was applied. The system was trained on the UCI SMS Spam Collection dataset, undergoing a rigorous five-stage "
    "text preprocessing pipeline including noise removal, tokenization, and stopword elimination. The developed system was deployed "
    "as a decoupled full-stack web application, featuring a React.js frontend dashboard and a FastAPI Python backend. Evaluation "
    "results demonstrate that the SMOTE-enhanced model achieved an accuracy of 96.32%, a precision of 89.15%, and a recall rate "
    "of 95.30%, proving highly effective at identifying disguised threat vectors while minimizing false positives. The deployed "
    "system provides a real-time, scalable solution for mobile network operators and end-users to intercept and mitigate SMS-based "
    "cyber threats efficiently."
)

body = doc.add_paragraph(abstract_text)
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
for run in body.runs:
    run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\5_Abstract.docx')
print('Abstract page created!')
