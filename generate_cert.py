from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph('CERTIFICATION')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(12)

# Body paragraph
body = doc.add_paragraph()
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
body.paragraph_format.space_after = Pt(24)
body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

r1 = body.add_run('This is to certify that the project titled "Design and Implementation of an Intelligent SMS Spam Detection System", submitted by ')
r1.font.size = Pt(12)

r2 = body.add_run('OGUNDIRAN CALEB AYODEJI (Matric No. 22/9789)')
r2.font.bold = True
r2.font.size = Pt(12)

r3 = body.add_run(', was duly supervised and has been approved as meeting part of the requirements for the award of the Bachelor of Science (B.Sc.) Degree in Cyber Security, Caleb University, Imota, Lagos State, Nigeria.')
r3.font.size = Pt(12)

# Helper to add a signature row using a table
def add_signature_row(doc, name, role):
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    # Remove borders
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Row 0: dotted lines
    sig_cell = table.cell(0, 0)
    sig_cell.text = '......................................'
    sig_cell.paragraphs[0].runs[0].font.size = Pt(12)

    date_cell = table.cell(0, 1)
    date_cell.text = '......................................'
    date_cell.paragraphs[0].runs[0].font.size = Pt(12)

    # Row 1: name | "Date."
    name_cell = table.cell(1, 0)
    name_cell.text = name
    name_cell.paragraphs[0].runs[0].font.size = Pt(12)

    date_label = table.cell(1, 1)
    date_label.text = 'Date.'
    date_label.paragraphs[0].runs[0].font.size = Pt(12)

    # Row 2: role
    role_cell = table.cell(2, 0)
    role_cell.text = role
    role_cell.paragraphs[0].runs[0].font.size = Pt(12)

    table.cell(2, 1).text = ''

    doc.add_paragraph()  # spacing after table

add_signature_row(doc, 'Prof. Aregbesola', '(Project Supervisor)')
add_signature_row(doc, 'Dr. Adegunwa Olajide', '(Head of Department)')
add_signature_row(doc, 'Prof. Moses Aregbesola', '(Dean of COCIS)')

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\2_Certification_Final.docx')
print('Certification updated with proper format!')
