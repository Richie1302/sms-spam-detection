from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(2.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_paragraph('DEDICATION')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.bold = True
    run.font.size = Pt(12)
title.paragraph_format.space_after = Pt(24)

# Dedication body - centered
body = doc.add_paragraph(
    'This project is dedicated to God Almighty, the Author and Finisher of my faith, '
    'the Source of all wisdom and strength. It is by His grace alone that I have come '
    'this far. To Him be all the glory, honour, and praise.'
)
body.alignment = WD_ALIGN_PARAGRAPH.CENTER
body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
for run in body.runs:
    run.font.size = Pt(12)

doc.save(r'C:\Users\𝕄ℝ ℕ𝕆\Desktop\SPAM DETECTION\docx\3_Dedication.docx')
print('Dedication page created!')
